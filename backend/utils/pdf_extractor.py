"""
utils/pdf_extractor.py
Extracts plain text from PDF, DOCX, or TXT files.
"""
import io

def extract_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return _from_docx(file_bytes)
    else:
        return file_bytes.decode("utf-8", errors="ignore")

def _from_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    pages = [page.get_text("text") for page in doc]
    doc.close()
    lines = [l.strip() for block in pages for l in block.splitlines() if l.strip()]
    return "\n".join(lines)

def _from_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)
